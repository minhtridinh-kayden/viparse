"""Tests for the DOCX extraction adapter."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

from viparse.detect import CONTENT_TYPE_DOCX
from viparse.engines.docx import DocxEngine
from viparse.errors import MissingDependency
from viparse.options import LoadOptions
from viparse.registry import EngineRegistry

docx = pytest.importorskip("docx")  # python-docx; skipped if the office extra is absent


def _make_docx(path: Path) -> Path:
    document = docx.Document()
    document.add_heading("Tiêu đề", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Tiếng Việt")
    run.font.name = ".VnTime"
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    document.add_paragraph("cuối")
    document.save(str(path))
    return path


def test_extract_returns_raw_extraction(tmp_path: Path) -> None:
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    assert raw.engine == "docx"
    assert raw.content_type == CONTENT_TYPE_DOCX
    assert "Tiếng Việt" in raw.text
    assert "cuối" in raw.text


def test_extract_captures_run_font_signal(tmp_path: Path) -> None:
    """The MVP acceptance: a run's font name is surfaced for the S3 detector."""
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    assert ".VnTime" in raw.signals["fonts"]


def test_extract_preserves_block_order(tmp_path: Path) -> None:
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    kinds = [block["type"] for block in raw.signals["blocks"]]
    assert kinds == ["heading", "paragraph", "table", "paragraph"]
    heading = raw.signals["blocks"][0]
    assert heading["level"] == 1
    assert heading["text"] == "Tiêu đề"


def test_extract_table_structure(tmp_path: Path) -> None:
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    table = next(block for block in raw.signals["blocks"] if block["type"] == "table")
    assert table["rows"] == [["A", "B"], ["C", "D"]]
    assert "A\tB" in raw.text


def test_empty_paragraphs_are_skipped(tmp_path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("")  # blank spacer
    document.add_paragraph("content")
    path = tmp_path / "b.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert [block["text"] for block in raw.signals["blocks"]] == ["content"]


def test_captures_style_inherited_font(tmp_path: Path) -> None:
    """A legacy font set via a character style (not per-run) is still surfaced."""
    from docx.enum.style import WD_STYLE_TYPE

    document = docx.Document()
    style = document.styles.add_style("Legacy", WD_STYLE_TYPE.CHARACTER)
    style.font.name = "VNI-Times"
    run = document.add_paragraph().add_run("styled text")
    run.style = document.styles["Legacy"]
    assert run.font.name is None  # font is only on the style, not the run
    path = tmp_path / "styled.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "VNI-Times" in raw.signals["fonts"]


def test_captures_paragraph_style_font(tmp_path: Path) -> None:
    """A font set on the paragraph style (not per-run) is surfaced too."""
    from docx.enum.style import WD_STYLE_TYPE

    document = docx.Document()
    style = document.styles.add_style("LegacyPara", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "VNI-Para"
    document.add_paragraph("body", style="LegacyPara")
    path = tmp_path / "parastyle.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "VNI-Para" in raw.signals["fonts"]


def test_style_font_helper_handles_missing_style() -> None:
    from viparse.engines.docx import _style_font

    assert _style_font(None) is None


def test_empty_content_control_is_ignored(tmp_path: Path) -> None:
    """A malformed w:sdt with no sdtContent is skipped without crashing."""
    from docx.oxml import OxmlElement

    document = docx.Document()
    document.add_paragraph("real content")
    document.element.body.append(OxmlElement("w:sdt"))  # no sdtContent child
    path = tmp_path / "emptysdt.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "real content" in raw.text


def test_captures_table_cell_font(tmp_path: Path) -> None:
    document = docx.Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    cell.paragraphs[0].add_run("cell").font.name = ".VnArial"
    path = tmp_path / "cellfont.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert ".VnArial" in raw.signals["fonts"]


def test_horizontally_merged_cells_are_not_duplicated(tmp_path: Path) -> None:
    document = docx.Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "merged"
    table.cell(0, 2).text = "third"
    table.cell(0, 0).merge(table.cell(0, 1))
    path = tmp_path / "merged.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    table_block = next(b for b in raw.signals["blocks"] if b["type"] == "table")
    assert table_block["rows"] == [["merged", "third"]]


def test_content_control_paragraphs_are_not_dropped(tmp_path: Path) -> None:
    """Text inside a w:sdt content control is extracted, not silently lost."""
    from docx.oxml import OxmlElement

    document = docx.Document()
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    para = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "inside control"
    run.append(text)
    para.append(run)
    content.append(para)
    sdt.append(content)
    document.element.body.append(sdt)
    path = tmp_path / "sdt.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "inside control" in raw.text


def test_supports_only_docx() -> None:
    engine = DocxEngine()
    assert engine.supports(CONTENT_TYPE_DOCX)
    assert not engine.supports("application/pdf")


def test_registry_selects_docx_engine() -> None:
    reg = EngineRegistry()
    reg.register(DocxEngine())
    assert isinstance(reg.select(CONTENT_TYPE_DOCX), DocxEngine)


def test_missing_dependency_raises_clear_error(monkeypatch: pytest.MonkeyPatch) -> None:
    """Without python-docx, extraction fails with actionable install guidance."""
    monkeypatch.setitem(sys.modules, "docx", None)
    with pytest.raises(MissingDependency, match=r"viparse\[office\]"):
        DocxEngine().extract("missing.docx", LoadOptions())
